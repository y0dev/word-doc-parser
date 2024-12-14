import os
from docx import Document
from lxml import etree
import utils.common_utils as comm_utils
from utils.time_to_read import TimeToRead
from .tags import Tags
from .image import Image

class WordDocParser:
    """
    This class parses a Word document (.docx) and extracts specific data.
    """

    def __init__(self, file_path, output_dir):
        """
        Initializes the WordDocParser object with the file path.

        Args:
            file_path (str): The path to the Word document (.docx) file.

        Raises:
            PermissionError: If the file cannot be opened due to permission issues.
        """
        self.file_path = file_path
        file_name, _ = os.path.splitext(os.path.basename(file_path)) 
        self.output_dir = os.path.join(output_dir, file_name)
        os.makedirs(self.output_dir.lower() , exist_ok=True)  # Create the output directory if it doesn't exist
        try:
            self.document = Document(file_path)
        except PermissionError:
            raise PermissionError(f"Cannot open the file {file_path}. Check read-only permissions.")
        self.__desc_start = False
        self.__word_count = 0
        self.data = {
            "metadata": {"id":"","type":"","title":"", "description":""},
            "headings": [],
            "paragraphs": [],
            "images": []
        }

    def extract_headings(self):
        """
        Extracts all headings from the document based on their paragraph style
        and stores them in a hierarchical structure with nested paragraphs.

        Iterates through each paragraph. If the paragraph style indicates a heading,
        a new dictionary entry is created with heading text, level, and an empty
        "paragraphs" list to hold nested content. Otherwise, it extracts formatted
        phrases and lists within the paragraph and adds them to a separate dictionary
        representing the paragraph data. This dictionary is then appended either to
        the current heading's "paragraphs" list (if inside a heading) or to the main
        "paragraphs" list in the `data` dictionary.
        """
        current_heading = None
        for paragraph in self.document.paragraphs:
            if not paragraph.text: continue
            
            if self.__update_metadata(paragraph):
                continue

            if paragraph.style.name.startswith('Heading'):
                current_heading = {
                    "text": paragraph.text.strip(),
                    "level": paragraph.style.name,
                    "paragraphs": []
                }
                self.data["headings"].append(current_heading)
            else:
                paragraph_data = {
                    "text": paragraph.text.strip(),
                    "bold_phrases": [],
                    "italic_phrases": [],
                    "underlined_phrases": [],
                    "lists": [],
                    "images": [],
                    "links": [],
                }
                self.__extract_formatted_phrases(paragraph, paragraph_data)
                self.__extract_links(paragraph, paragraph_data)
                found_list = self.extract_lists(paragraph, paragraph_data)
                found_image = self.__extract_images_from_para(paragraph, paragraph_data)

                if not found_list and not found_image:
                    self.__word_count += len(paragraph.text.strip())
                    if current_heading:
                        current_heading["paragraphs"].append(paragraph_data)
                    else:
                        self.data["paragraphs"].append(paragraph_data)
        time_to_read = TimeToRead(self.__word_count)
        self.data["metadata"]["time_to_read"] = time_to_read.get_time_as_obj()

    def __extract_formatted_phrases(self, paragraph, paragraph_data):
        """
        Extracts phrases that are entirely bold, italic, or underlined within a paragraph.

        Iterates through each text run (formatted text segment) in the paragraph.
        Maintains temporary lists for ongoing bold, italic, and underlined phrases.
        It only adds a phrase to the corresponding list in the `paragraph_data` dictionary
        ("bold_phrases", "italic_phrases", or "underlined_phrases") when the phrase ends
        (no longer bold/italic/underlined). This ensures that only complete formatted phrases are captured.

        Args:
            paragraph (docx.paragraph.Paragraph): A paragraph object from the Word document.
            paragraph_data (dict): A dictionary to store extracted data for the current paragraph.
        """
        bold_phrase, italic_phrase, underlined_phrase = [], [], []

        for run in paragraph.runs:
            word = run.text.strip()

            if run.bold:
                bold_phrase.append(word)
            elif bold_phrase:
                paragraph_data["bold_phrases"].append(" ".join(bold_phrase))
                bold_phrase = []

            if run.italic:
                italic_phrase.append(word)
            elif italic_phrase:
                paragraph_data["italic_phrases"].append(" ".join(italic_phrase))
                italic_phrase = []

            if run.underline:
                underlined_phrase.append(word)
            elif underlined_phrase:
                paragraph_data["underlined_phrases"].append(" ".join(underlined_phrase))
                underlined_phrase = []

        if bold_phrase:
            paragraph_data["bold_phrases"].append(" ".join(bold_phrase))
        if italic_phrase:
            paragraph_data["italic_phrases"].append(" ".join(italic_phrase))
        if underlined_phrase:
            paragraph_data["underlined_phrases"].append(" ".join(underlined_phrase))
    
    def __extract_links(self, paragraph, paragraph_data):
        """
        Extracts hyperlinks within a paragraph.

        Checks if the paragraph contains any hyperlinks and extracts their text and target URL.
        
        Args:
            paragraph (docx.paragraph.Paragraph): A paragraph object from the Word document.
            paragraph_data (dict): A dictionary to store extracted data for the current paragraph.
        """
        # Initialize an empty list to store links
        links = []
        xml_content = paragraph._element
        for rel in xml_content.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink"):
            # Extract relationship ID
            link_id = rel.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id")
            if link_id and link_id in self.document.part.rels:
                # Get link target and text
                link_target = self.document.part.rels[link_id]._target
                link_text = "".join(t.text for t in rel.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"))
                links.append({
                    "text": link_text,
                    "target": link_target
                })

        # Store the extracted links in the paragraph data dictionary
        if links:
            paragraph_data['links'] = links


    def extract_lists(self, paragraph, paragraph_data):
        """
        Extracts numbered and bulleted lists within a paragraph.

        Checks if the paragraph style indicates a list ("List Paragraph"). If so,
        it extracts the list text and its indent level. The function primarily focuses
        on retrieving the last list, considering sublists based on indentation.

        Args:
            paragraph (docx.paragraph.Paragraph): A paragraph object from the Word document.
            paragraph_data (dict): A dictionary to store extracted data for the current paragraph.
        """
        if paragraph.style.name in ["List Paragraph"]:
            list_text = paragraph.text.strip()
            indent_level = paragraph.paragraph_format.left_indent.pt if paragraph.paragraph_format.left_indent else 0

            last_list_container = None
            if len(self.data["headings"]) > 0:
                if len(self.data["headings"][-1]["paragraphs"]) > 0:
                    last_list_container = self.data["headings"][-1]["paragraphs"][-1]
            elif len(self.data["paragraphs"]) > 0:
                last_list_container = self.data["paragraphs"][-1]

            if last_list_container:
                last_list = last_list_container["lists"][-1] if len(last_list_container["lists"]) > 0 else None
                if last_list:
                    if indent_level > last_list["indent_level"]:
                        if "sublist" not in last_list:
                            last_list["sublist"] = []
                        last_list["sublist"].append({
                            "text": list_text,
                            "indent_level": indent_level
                        })
                    else:
                        last_list_container["lists"].append({
                            "text": list_text,
                            "indent_level": indent_level
                        })
                else:
                    last_list_container["lists"].append({
                        "text": list_text,
                        "indent_level": indent_level
                    })
            else:
                paragraph_data["lists"].append({
                    "text": list_text,
                    "indent_level": indent_level
                })
            return True
        return False
    def __extract_images_from_para(self, paragraph, paragraph_data):

        last_list_container = None
        if len(self.data["headings"]) > 0:
            if len(self.data["headings"][-1]["paragraphs"]) > 0:
                last_list_container = self.data["headings"][-1]["paragraphs"][-1]
        elif len(self.data["paragraphs"]) > 0:
            last_list_container = self.data["paragraphs"][-1]

        for image in self.data["images"]:
            if image["caption"] == paragraph.text.strip():
                if last_list_container:
                    last_list_container["images"].append({
                        "id": image["id"],
                        "caption": image["caption"]
                    })
                else:
                    paragraph_data["images"].append({
                        "id": image["id"],
                        "caption": image["caption"]
                    })
                return True
        return False


    def __extract_images(self):
        """
        Extracts images and their captions from the Word document and adds them to the data structure.

        This function iterates through inline shapes (embedded images) within the document.
        For each image, it:
            1. Extracts the image data from the relationship part.
            2. Saves the image data to a file with a descriptive filename (e.g., extracted_image_1.jpg).
            3. Attempts to find a corresponding caption:
                - Checks paragraphs following the image's location in the document.
                - If the next paragraph has the "caption" style, extracts its text as the caption.
            4. Adds an entry to the "images" list within the `data` dictionary.
                - The entry includes:
                    - "id": A unique identifier for the image, formatted with leading zeros using `comm_utils.fill_string_with_zeros`.
                    - "data": The extracted image data.
                    - "caption": The extracted caption text (if available).

        Args:
            self: An instance of the class responsible for document processing.
        """

        # Create output directory for extracted images (if it doesn't exist)
        output_dir = os.path.join(self.output_dir, "images")
        os.makedirs(output_dir, exist_ok=True)

        # Extract embedded images
        for i, shape in enumerate(self.document.inline_shapes):
            # Get image data and related part
            image_data = shape._inline.graphic.graphicData.pic.blipFill.blip.embed
            image_part = self.document.part.related_parts[image_data]

            # Save image to a file
            image_filename = os.path.join(output_dir, f"extracted_image_{i + 1}.jpg")
            with open(image_filename, "wb") as f:
                f.write(image_part.blob)

            # Find corresponding caption (if any)
            caption = ""
            found_image = False
            for para in self.document.paragraphs:
                if image_data in para._element.xml:  # Check if image data is present in paragraph XML
                    found_image = True
                elif found_image:  # Check next paragraph if image data found previously
                    if para.style.name.lower() == "caption":
                        caption = para.text.strip()  # Extract caption text if paragraph style is "caption"
                    break

            # Print confirmation and add image data to output
            print(f"Saved image: {image_filename}")
            self.data["images"].append({
                "id": f"{comm_utils.fill_string_with_zeros(i+1,3)}",  # Create unique ID with leading zeros
                "data": image_data,
                "caption": caption
            })

    def __extract_code(self, paragraph):
        pass

    def __update_metadata(self, paragraph):
        """
        Private helper function to update the metadata dictionary based on specific keywords.

        This function avoids code duplication and improves readability by encapsulating the logic
        for updating metadata based on keywords within paragraphs.

        Args:
            paragraph (docx.paragraph.Paragraph): A paragraph object from the Word document.
        """
        text = paragraph.text.lower().strip()
        if self.__desc_start:
            self.data["metadata"]["description"] = text
            self.__desc_start = False
            return True

        if text.startswith("article-id"):
            self.data["metadata"]["id"] = text.split("=")[1].strip().replace(" ","-")
            return True
        elif text.startswith("article-type"):
            tags = Tags(text.split("=")[1].strip())
            image = Image(text.split("=")[1].strip())
            self.data["metadata"]["type"] = text.split("=")[1].strip()
            self.data["metadata"]["tags"] = tags.get_tag_list()
            self.data["metadata"]["image"] = image.get_image()
            return True
        elif text.startswith("article-title"):
            self.data["metadata"]["title"] = text.split("=")[1].strip()
            return True
        elif text.startswith("description"):
            self.__desc_start = True
            return True
        return False
    
    def parse_document(self):
        """
        The main method to initiate the parsing process.

        Calls all the individual extraction methods like `extract_headings`,
        `extract_formatted_phrases`, and `extract_lists` to extract the desired data.
        Returns the extracted data stored in the `data` dictionary.

        Returns:
            dict: A dictionary containing extracted information from the document.
        """
        self.__extract_images()
        self.extract_headings()
        timestamp = None
        if comm_utils.confirm_update("Do you want to update the list of choices?"):
            timestamp = comm_utils.get_user_date() 
        else:
            timestamp = comm_utils.get_file_creation_time(self.file_path)
        self.data["metadata"]["date"] = str(comm_utils.generate_timestamp_millis(timestamp))

        return self.data
