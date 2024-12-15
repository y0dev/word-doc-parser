import unittest
from lib.word_parser.word_doc_parser import WordDocParser
from docx import Document
import os

class TestWordDocParser(unittest.TestCase):
    def setUp(self):
        """ Create a temporary DOCX file for testing """
        self.test_file = "tests/test_doc.docx"
        self.output_dir = os.path.join("tests","output")
        os.makedirs(self.output_dir.lower() , exist_ok=True)
        doc = Document()
        doc.add_heading("Test Heading 1", level=1)
        doc.add_paragraph("This is a normal paragraph.")
        p = doc.add_paragraph()
        p.add_run("Bold Text").bold = True
        p.add_run(" Normal Text")
        p.add_run(" Italic Text").italic = True
        doc.save(self.test_file)

    def tearDown(self):
        """ Clean up the temporary file after testing """
        if os.path.exists(self.test_file):
            os.remove(self.test_file)

    def test_extract_headings(self):
        parser = WordDocParser(self.test_file,self.output_dir)
        parser.extract_headings()
        self.assertEqual(len(parser.data["headings"]), 1)
        self.assertEqual(parser.data["headings"][0]["text"], "Test Heading 1")

    def test_extract_formatted_phrases(self):
        parser = WordDocParser(self.test_file,self.output_dir)

        # self.assertIn("Bold Text", parser.data["bold_phrases"])
        # self.assertIn("Italic Text", parser.data["italic_phrases"])
        # self.assertNotIn("Normal Text", parser.data["bold_phrases"])
        # self.assertNotIn("Normal Text", parser.data["italic_phrases"])

if __name__ == '__main__':
    unittest.main()
